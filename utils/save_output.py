from docx import Document
from docx2pdf import convert

def save_proposal_to_docx(proposal_text: str, filename: str = "polished_proposal.docx"):
    doc = Document()
    doc.add_heading("Proposal", 0)
    for line in proposal_text.strip().split("\n"):
        if line.startswith("**") and line.endswith("**"):
            doc.add_heading(line.replace("**", ""), level=1)
        elif line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(filename)
    return filename

def convert_docx_to_pdf(docx_file, pdf_file="polished_proposal.pdf"):
    convert(docx_file, pdf_file)
